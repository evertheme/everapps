"""
Generates frontend/public/templates/requirements-template.docx

Run from the repo root:
    python scripts/generate_requirements_template.py

Requires:  pip install python-docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = Path(__file__).parent.parent / "frontend" / "public" / "templates" / "requirements-template.docx"


# ── Colour palette ─────────────────────────────────────────────────────────────
BRAND_BLUE   = RGBColor(0x26, 0x63, 0xEB)   # brand-600
BRAND_LIGHT  = RGBColor(0xEB, 0xF2, 0xFF)   # brand-50
DARK_TEXT    = RGBColor(0x11, 0x18, 0x27)   # gray-900
MID_GREY     = RGBColor(0x6B, 0x72, 0x80)   # gray-500
LIGHT_GREY   = RGBColor(0xF9, 0xFA, 0xFB)   # gray-50
BORDER_GREY  = RGBColor(0xE5, 0xE7, 0xEB)   # gray-200
GREEN_TEXT   = RGBColor(0x05, 0x96, 0x69)   # emerald-600
AMBER_TEXT   = RGBColor(0xD9, 0x77, 0x06)   # amber-600


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def set_cell_border(cell, sides=("top", "bottom", "left", "right"), color="E5E7EB", sz=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = BRAND_BLUE
        run.font.size      = Pt(14)
        run.bold           = True
    elif level == 2:
        run.font.color.rgb = DARK_TEXT
        run.font.size      = Pt(12)
        run.bold           = True
    else:
        run.font.color.rgb = MID_GREY
        run.font.size      = Pt(11)
        run.bold           = True
    return p


def add_guidance(doc: Document, text: str):
    """Italic grey guidance note."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.italic          = True
    run.font.color.rgb  = MID_GREY
    run.font.size       = Pt(9.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_placeholder(doc: Document, text: str, indent: bool = False):
    """Regular placeholder paragraph."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)   # gray-400
    run.font.size      = Pt(10)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None):
    """Styled table with a branded header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold           = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.size      = Pt(9)
        hdr_cells[i].paragraphs[0].alignment              = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_bg(hdr_cells[i], BRAND_BLUE)

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg    = LIGHT_GREY if r_idx % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            cells[c_idx].paragraphs[0].runs[0].font.size      = Pt(9)
            cells[c_idx].paragraphs[0].runs[0].font.color.rgb = DARK_TEXT
            set_cell_bg(cells[c_idx], bg)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return table


def add_divider(doc: Document):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "E5E7EB")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)


# ── Document setup ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default paragraph font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ── Cover / title ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("[Project Name]")
    tr.bold           = True
    tr.font.size      = Pt(24)
    tr.font.color.rgb = BRAND_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("Software Requirements Document")
    sr.font.size      = Pt(14)
    sr.font.color.rgb = MID_GREY

    doc.add_paragraph()

    add_table(doc,
        headers=["Field", "Value"],
        rows=[
            ["Document Title",   "[Project Name] Software Requirements Specification"],
            ["Version",          "0.1"],
            ["Date",             "[YYYY-MM-DD]"],
            ["Author(s)",        "[Full Name, Role]"],
            ["Approval Status",  "Draft"],
            ["Project Description", "[One sentence describing what this project builds and for whom.]"],
        ],
        col_widths=[2.0, 4.5],
    )

    add_guidance(doc,
        "How to use this template: replace every [placeholder] with your project's details. "
        "Delete guidance notes (italic grey text) before uploading to everapps. "
        "The more completely you fill this in, the higher the quality of your generated user story backlog."
    )

    add_divider(doc)
    doc.add_page_break()

    # ── Section 1: Executive Summary ──────────────────────────────────────────
    add_heading(doc, "1. Executive Summary")
    add_guidance(doc, "Write 1–3 paragraphs for a non-technical stakeholder. Cover: the problem, who it is for, and the intended outcome.")
    add_placeholder(doc, "[Problem statement — what pain or inefficiency exists today, and why it matters?]")
    doc.add_paragraph()
    add_placeholder(doc, "[Solution — what is being built, for whom, and how does it address the problem?]")
    doc.add_paragraph()
    add_placeholder(doc, "[Outcome — what does success look like once this is delivered?]")
    add_divider(doc)

    # ── Section 2: Project Context & Business Objectives ──────────────────────
    add_heading(doc, "2. Project Context & Business Objectives")

    add_heading(doc, "Business Problem Statement", level=2)
    add_placeholder(doc, "[Describe the core business problem or opportunity in plain language.]")

    add_heading(doc, "Business Objectives", level=2)
    add_guidance(doc, "List 3–7 measurable objectives. Use 'to [verb] [metric] by [target]' format where possible.")
    for i, obj in enumerate([
        "To [achieve outcome] by [measurable target] within [timeframe].",
        "To [reduce / increase / improve] [metric] from [current state] to [desired state].",
        "To [enable / automate / eliminate] [activity] for [user group].",
        "[Add more objectives as needed]",
    ], 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(obj).font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        p.runs[-1].font.size = Pt(10)

    add_heading(doc, "Current State vs. Desired Future State", level=2)
    add_table(doc,
        headers=["Aspect", "Current State", "Desired Future State"],
        rows=[
            ["[Process / System]", "[How it works today]", "[How it should work]"],
            ["[Data / Reporting]", "[Current limitations]", "[Target capability]"],
            ["[User Experience]",  "[Pain points]",         "[Improved experience]"],
        ],
        col_widths=[1.8, 2.6, 2.6],
    )

    add_heading(doc, "Organisational Alignment", level=2)
    add_placeholder(doc, "[Explain how this project supports the organisation's broader strategy, OKRs, or roadmap.]")
    add_divider(doc)

    # ── Section 3: Scope ──────────────────────────────────────────────────────
    add_heading(doc, "3. Scope")

    add_heading(doc, "In Scope", level=2)
    add_guidance(doc, "List features, user groups, and system boundaries explicitly included.")
    for item in [
        "[Feature area 1 — brief description]",
        "[Feature area 2 — brief description]",
        "[User group 1] can [do X, Y, Z]",
        "Integration with [System A] for [purpose]",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        p.runs[-1].font.size = Pt(10)

    add_heading(doc, "Out of Scope", level=2)
    add_guidance(doc, "Explicit exclusions prevent scope creep. Be specific.")
    for item in [
        "[Feature or capability] — deferred to Phase 2",
        "[User group or workflow] — not addressed in this release",
        "[Integration with System B] — future enhancement",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        p.runs[-1].font.size = Pt(10)

    add_heading(doc, "Integration Touchpoints", level=2)
    add_table(doc,
        headers=["System", "Direction", "Purpose"],
        rows=[
            ["[System / API name]", "Inbound / Outbound / Bidirectional", "[What data is exchanged and why]"],
            ["[System / API name]", "Inbound / Outbound / Bidirectional", "[What data is exchanged and why]"],
        ],
        col_widths=[1.8, 2.0, 3.2],
    )
    add_divider(doc)

    # ── Section 4: Stakeholders & User Personas ───────────────────────────────
    add_heading(doc, "4. Stakeholders & User Personas")

    add_heading(doc, "Stakeholders", level=2)
    add_table(doc,
        headers=["Name / Role", "Organisation", "Interest / Influence", "Key Concerns"],
        rows=[
            ["[Executive Sponsor]",       "[Dept]", "High / High",   "ROI, delivery timeline"],
            ["[Product Owner]",           "[Dept]", "High / High",   "Scope, priorities"],
            ["[IT / Engineering Lead]",   "[Dept]", "Medium / High", "Architecture, security"],
            ["[End-User Representative]", "[Dept]", "High / Low",    "Usability, workflow"],
        ],
        col_widths=[1.8, 1.4, 1.6, 2.2],
    )

    add_heading(doc, "User Personas", level=2)
    add_guidance(doc, "For each persona: who they are, goals, pain points, and how they interact with the system.")
    for i in range(1, 4):
        add_heading(doc, f"Persona {i}: [Role Title, e.g. Operations Manager]", level=3)
        add_table(doc,
            headers=["Attribute", "Detail"],
            rows=[
                ["Who",          "[Brief description — seniority, team, technical proficiency]"],
                ["Goals",        "[What they are trying to achieve with this system]"],
                ["Pain Points",  "[What frustrates them about the current state]"],
                ["Interactions", "[Which parts of the system they use and how often]"],
            ],
            col_widths=[1.5, 5.5],
        )
    add_divider(doc)

    # ── Section 5: Functional Requirements ───────────────────────────────────
    add_heading(doc, "5. Functional Requirements")
    add_guidance(doc,
        "Each requirement needs: a unique ID (FR-NNN), 'The system shall…' language, "
        "a MoSCoW priority (Must-have / Should-have / Nice-to-have), and acceptance criteria "
        "in Given / When / Then format."
    )

    for area_num, area_name, reqs in [
        (1, "[Feature Area — e.g. User Authentication]", [
            ("FR-001", "The system shall allow users to register with an email address and password.", "Must-have",
             "Given a new user submits a valid email and password, when they click Register, then an account is created and they are redirected to the dashboard."),
            ("FR-002", "The system shall send a verification email upon registration.", "Must-have",
             "Given a user has registered, when the account is created, then a verification email is sent within 60 seconds."),
            ("FR-003", "The system shall allow users to reset their password via email.", "Must-have",
             "Given a user requests a password reset, when they submit their email, then a reset link is sent within 60 seconds."),
        ]),
        (2, "[Feature Area — e.g. Dashboard & Reporting]", [
            ("FR-010", "The system shall display a summary dashboard upon login.", "Must-have",
             "Given a logged-in user navigates to the dashboard, when the page loads, then key metrics are displayed within 3 seconds."),
            ("FR-011", "The system shall allow users to export reports as CSV.", "Should-have",
             "Given a user views a report, when they click Export, then a CSV file is downloaded within 10 seconds."),
        ]),
        (3, "[Feature Area — add more as needed]", [
            ("FR-020", "The system shall [describe what the system must do].", "Must-have",
             "Given [context], when [action], then [expected outcome]."),
        ]),
    ]:
        add_heading(doc, f"5.{area_num} {area_name}", level=2)
        add_table(doc,
            headers=["ID", "Requirement", "Priority"],
            rows=[[r[0], r[1], r[2]] for r in reqs],
            col_widths=[0.8, 5.2, 1.0],
        )
        add_heading(doc, "Acceptance Criteria", level=3)
        for r in reqs:
            p = doc.add_paragraph(style="List Bullet")
            bold_run = p.add_run(f"{r[0]}: ")
            bold_run.bold = True
            bold_run.font.size = Pt(9.5)
            body_run = p.add_run(r[3])
            body_run.font.size = Pt(9.5)
            body_run.font.color.rgb = DARK_TEXT
        doc.add_paragraph()

    add_divider(doc)

    # ── Section 6: Non-Functional Requirements ────────────────────────────────
    add_heading(doc, "6. Non-Functional Requirements")

    add_heading(doc, "Performance", level=2)
    add_table(doc,
        headers=["ID", "Requirement", "Target"],
        rows=[
            ["NFR-PERF-001", "Page load time (95th percentile)",    "< 3 seconds"],
            ["NFR-PERF-002", "API response time (95th percentile)", "< 500 ms"],
            ["NFR-PERF-003", "Concurrent users supported",          "[N] simultaneous users"],
        ],
        col_widths=[1.4, 3.8, 1.8],
    )

    add_heading(doc, "Security", level=2)
    add_table(doc,
        headers=["ID", "Requirement", "Detail"],
        rows=[
            ["NFR-SEC-001", "Authentication",   "Multi-factor authentication (MFA) required for [role(s)]"],
            ["NFR-SEC-002", "Authorisation",    "Role-based access control (RBAC) with [N] roles"],
            ["NFR-SEC-003", "Data encryption",  "AES-256 at rest; TLS 1.2+ in transit"],
            ["NFR-SEC-004", "Audit logging",    "All user / admin actions logged with timestamp and user ID"],
            ["NFR-SEC-005", "Compliance",       "[GDPR / SOC 2 / HIPAA / PCI-DSS — delete as applicable]"],
        ],
        col_widths=[1.4, 1.8, 3.8],
    )

    add_heading(doc, "Scalability", level=2)
    add_table(doc,
        headers=["ID", "Requirement", "Detail"],
        rows=[
            ["NFR-SCAL-001", "Horizontal scaling",  "Application tier must support auto-scaling under load"],
            ["NFR-SCAL-002", "Data volume",          "Must support up to [N] records / [X] GB"],
            ["NFR-SCAL-003", "Growth projection",    "Architecture must support 3× volume within 2 years without redesign"],
        ],
        col_widths=[1.4, 1.8, 3.8],
    )

    add_heading(doc, "Reliability", level=2)
    add_table(doc,
        headers=["ID", "Requirement", "Target"],
        rows=[
            ["NFR-REL-001", "Uptime SLA",                      "99.9% availability"],
            ["NFR-REL-002", "Recovery Time Objective (RTO)",   "< [N] hours"],
            ["NFR-REL-003", "Recovery Point Objective (RPO)",  "< [N] hours of data loss"],
            ["NFR-REL-004", "Backup frequency",                "Daily automated backups retained [N] days"],
        ],
        col_widths=[1.4, 2.8, 2.8],
    )

    add_heading(doc, "Usability & Accessibility", level=2)
    add_table(doc,
        headers=["ID", "Requirement", "Detail"],
        rows=[
            ["NFR-USE-001", "Accessibility standard", "WCAG 2.1 Level AA compliance"],
            ["NFR-USE-002", "Supported browsers",     "Latest 2 versions of Chrome, Firefox, Safari, Edge"],
            ["NFR-USE-003", "Responsive design",      "Functional on desktop (1024px+) and tablet (768px+)"],
        ],
        col_widths=[1.4, 1.8, 3.8],
    )

    add_heading(doc, "Maintainability", level=2)
    add_table(doc,
        headers=["ID", "Requirement", "Detail"],
        rows=[
            ["NFR-MAIN-001", "Code standards",   "Coding standards enforced via linting"],
            ["NFR-MAIN-002", "Test coverage",    "Minimum [N]% unit test coverage on business logic"],
            ["NFR-MAIN-003", "Deployment",       "Zero-downtime deployments via CI/CD"],
            ["NFR-MAIN-004", "Documentation",    "API documented via OpenAPI 3.0"],
        ],
        col_widths=[1.4, 1.8, 3.8],
    )
    add_divider(doc)

    # ── Section 7: Data & Integration Requirements ────────────────────────────
    add_heading(doc, "7. Data & Integration Requirements")

    add_heading(doc, "Key Data Entities", level=2)
    add_table(doc,
        headers=["Entity", "Description", "Key Attributes", "Relationships"],
        rows=[
            ["[Entity 1, e.g. User]",  "[What it represents]", "id, email, role, created_at", "Has many [Orders]"],
            ["[Entity 2, e.g. Order]", "[What it represents]", "id, user_id, status, total",  "Belongs to [User]"],
            ["[Entity 3]",             "[What it represents]", "…",                            "…"],
        ],
        col_widths=[1.4, 1.8, 2.0, 1.8],
    )

    add_heading(doc, "Data Retention & Deletion", level=2)
    add_table(doc,
        headers=["Data Type", "Retention Period", "Deletion Policy"],
        rows=[
            ["[User account data]",    "[Duration of account + 30 days]",  "[Hard delete / anonymise] on closure"],
            ["[Transaction records]",  "[e.g. 7 years]",                    "Archived after [N] years"],
            ["[Log / audit data]",     "[e.g. 90 days]",                    "Auto-purged after retention period"],
        ],
        col_widths=[2.0, 2.0, 3.0],
    )

    add_heading(doc, "External Integrations", level=2)
    add_table(doc,
        headers=["System", "Protocol", "Data Exchanged", "Direction", "Frequency"],
        rows=[
            ["[System name]", "REST API / Webhook", "[What data]", "Inbound / Outbound", "Real-time / Batch"],
            ["[System name]", "REST API",           "[What data]", "Outbound",           "On-demand"],
        ],
        col_widths=[1.4, 1.2, 2.0, 1.2, 1.2],
    )

    add_heading(doc, "Data Migration", level=2)
    add_guidance(doc, "Complete this section only if existing data needs to be migrated from a legacy system.")
    add_table(doc,
        headers=["Attribute", "Detail"],
        rows=[
            ["Source system",         "[Name and description of existing system]"],
            ["Volume",                "Approximately [N] records across [M] entity types"],
            ["Migration approach",    "[One-time cutover / phased / parallel run]"],
            ["Data cleansing",        "[Yes — describe known quality issues / No]"],
            ["Rollback plan",         "[How to revert if migration fails]"],
        ],
        col_widths=[1.8, 5.2],
    )
    add_divider(doc)

    # ── Section 8: Constraints & Assumptions ──────────────────────────────────
    add_heading(doc, "8. Constraints & Assumptions")

    add_heading(doc, "Constraints", level=2)
    add_guidance(doc, "Constraints are fixed limitations that cannot be changed.")
    add_table(doc,
        headers=["ID", "Constraint", "Category", "Impact"],
        rows=[
            ["C-001", "[e.g. Must deploy on existing AWS infrastructure]",   "Technology",  "Limits cloud-provider choices"],
            ["C-002", "[e.g. Budget cap of $X for Year 1]",                  "Budget",      "Limits licensing and headcount"],
            ["C-003", "[e.g. Must go live by YYYY-MM-DD]",                   "Timeline",    "Forces scope prioritisation"],
            ["C-004", "[e.g. Must comply with GDPR]",                        "Regulatory",  "Data residency and consent flows required"],
        ],
        col_widths=[0.7, 3.0, 1.2, 2.1],
    )

    add_heading(doc, "Assumptions", level=2)
    add_guidance(doc, "If any assumption proves false, requirements may need to change. Flag these for stakeholders.")
    add_table(doc,
        headers=["ID", "Assumption", "Owner", "Risk if False"],
        rows=[
            ["A-001", "[e.g. Users have reliable internet access]",             "[Product]",     "Offline capability would be required"],
            ["A-002", "[e.g. Existing API from System X is stable]",           "[Engineering]", "Integration effort significantly increases"],
            ["A-003", "[e.g. Dedicated QA resource available from Month N]",    "[Delivery]",    "Test phase timeline at risk"],
        ],
        col_widths=[0.7, 3.0, 1.3, 2.0],
    )

    add_heading(doc, "Dependencies", level=2)
    add_table(doc,
        headers=["ID", "Dependency", "Owner", "Due Date"],
        rows=[
            ["D-001", "[e.g. Design system / brand guidelines finalised]", "[Design team]", "[Date]"],
            ["D-002", "[e.g. API credentials from third-party vendor]",     "[Vendor name]", "[Date]"],
            ["D-003", "[e.g. Infrastructure provisioning approved]",        "[IT / DevOps]", "[Date]"],
        ],
        col_widths=[0.7, 3.6, 1.5, 1.2],
    )
    add_divider(doc)

    # ── Section 9: Success Metrics & Acceptance Criteria ─────────────────────
    add_heading(doc, "9. Success Metrics & Acceptance Criteria")

    add_heading(doc, "Definition of Done (Project Level)", level=2)
    add_guidance(doc, "The project is considered complete when all of the following are true.")
    for item in [
        "All Must-have functional requirements pass user acceptance testing (UAT)",
        "All Non-Functional Requirements verified in staging environment",
        "Security penetration testing completed with no Critical or High findings open",
        "Accessibility audit passed (WCAG 2.1 AA)",
        "Go-live sign-off obtained from [Executive Sponsor / Product Owner]",
        "Operations runbook and support handover documentation delivered",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"☐  {item}").font.size = Pt(10)

    add_heading(doc, "Post-Launch KPIs", level=2)
    add_table(doc,
        headers=["KPI", "Baseline", "Target 30 days", "Target 90 days", "Measurement"],
        rows=[
            ["[Task completion time]", "[N minutes]",    "[< X minutes]",        "[< Y minutes]",        "In-app analytics"],
            ["[User adoption rate]",   "N/A",            "[N% of target users]", "[M% of target users]", "Login analytics"],
            ["[Support ticket rate]",  "[N / week]",     "[< X / week]",         "[< Y / week]",         "Support system"],
            ["[Process cost]",         "[$X per unit]",  "[< $Y]",               "[< $Z]",               "Finance reporting"],
        ],
        col_widths=[1.6, 1.0, 1.2, 1.2, 1.5],
    )

    add_heading(doc, "UAT Test Scenarios", level=2)
    add_table(doc,
        headers=["Test Scenario", "Persona", "Pass Criteria"],
        rows=[
            ["[Register and complete onboarding]",           "[New User]",   "All steps completed without errors in < 5 minutes"],
            ["[Submit a complete [entity]]",                 "[Persona 1]",  "Submission confirmed; appears in dashboard within 10 seconds"],
            ["[Export [report] as CSV]",                     "[Persona 2]",  "CSV downloaded; data matches on-screen data"],
            ["[Admin reviews and approves [item]]",          "[Admin]",      "Status updated; user notified within 60 seconds"],
        ],
        col_widths=[2.2, 1.2, 3.6],
    )
    add_divider(doc)

    # ── Section 10: Timeline & Prioritisation ─────────────────────────────────
    add_heading(doc, "10. Timeline & Prioritisation")
    add_guidance(doc, "This section is optional but helps story generation prioritise the backlog.")

    add_heading(doc, "High-Level Milestones", level=2)
    add_table(doc,
        headers=["Milestone", "Description", "Target Date"],
        rows=[
            ["Project Kick-off",        "Team onboarded, environments provisioned", "[Date]"],
            ["Design Complete",         "UI/UX designs approved",                   "[Date]"],
            ["MVP Development Complete","All Must-have requirements built",          "[Date]"],
            ["UAT Start",               "Stakeholder testing begins",                "[Date]"],
            ["UAT Sign-off",            "All UAT criteria passed",                   "[Date]"],
            ["Go-Live",                 "Production deployment",                     "[Date]"],
            ["Post-Launch Review",      "30-day KPI review",                         "[Date]"],
        ],
        col_widths=[2.0, 3.5, 1.5],
    )

    add_heading(doc, "MoSCoW Feature Prioritisation", level=2)
    add_table(doc,
        headers=["Priority", "Feature Areas"],
        rows=[
            ["Must-have",    "[Feature area 1], [Feature area 2], [Feature area 3]"],
            ["Should-have",  "[Feature area 4], [Feature area 5]"],
            ["Nice-to-have", "[Feature area 6], [Feature area 7]"],
            ["Out of scope", "[Feature area 8] — Phase 2"],
        ],
        col_widths=[1.5, 5.5],
    )

    add_heading(doc, "Release Phasing", level=2)
    for phase, desc in [
        ("MVP (Phase 1):", "[Brief description of what ships first and why.]"),
        ("Phase 2:",       "[What follows, and the dependency on Phase 1.]"),
        ("Phase 3 / Future:", "[Longer-term vision items.]"),
    ]:
        p = doc.add_paragraph()
        p.add_run(phase + " ").bold = True
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)
        p.add_run(desc).font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        p.runs[-1].font.size = Pt(10)

    add_divider(doc)

    # ── Section 11: Glossary ──────────────────────────────────────────────────
    add_heading(doc, "11. Glossary")
    add_guidance(doc, "Define domain-specific terms, acronyms, and abbreviations so all readers share a common vocabulary.")
    add_table(doc,
        headers=["Term", "Definition"],
        rows=[
            ["[Acronym / Term]", "[Plain-English definition]"],
            ["[Acronym / Term]", "[Plain-English definition]"],
            ["API",   "Application Programming Interface — a way for software systems to communicate"],
            ["RBAC",  "Role-Based Access Control — permissions assigned to roles rather than individuals"],
            ["SLA",   "Service Level Agreement — a commitment to a defined level of service availability"],
            ["UAT",   "User Acceptance Testing — stakeholder verification that the system meets requirements"],
            ["MoSCoW","Must-have / Should-have / Could-have / Won't-have — a prioritisation framework"],
            ["RTO",   "Recovery Time Objective — maximum acceptable downtime after an incident"],
            ["RPO",   "Recovery Point Objective — maximum acceptable data loss after an incident"],
        ],
        col_widths=[2.0, 5.0],
    )

    # ── Footer note ───────────────────────────────────────────────────────────
    add_divider(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("Generated by everapps — Requirements to Backlog, intelligently.")
    fr.italic          = True
    fr.font.size       = Pt(8.5)
    fr.font.color.rgb  = MID_GREY

    hint_p = doc.add_paragraph()
    hint_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hint_p.add_run("Upload this completed document to everapps to generate your user story backlog.")
    hr.italic         = True
    hr.font.size      = Pt(8.5)
    hr.font.color.rgb = MID_GREY

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Template written to {OUTPUT}")


if __name__ == "__main__":
    build()
